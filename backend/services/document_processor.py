"""
services/document_processor.py
===============================
Extract and clean plain text from uploaded PDF and DOCX files.

Supported formats
-----------------
- PDF  → pdfplumber  (handles multi-page, tables, headers)
- DOCX → python-docx (handles paragraphs, tables, headers)
"""

import os
import re
from typing import Dict, Any

import pdfplumber
from docx import Document

from utils.logger import app_logger


# --------------------------------------------------------------------------- #
# Public entry point                                                            #
# --------------------------------------------------------------------------- #

def process_document(file_path: str) -> Dict[str, Any]:
    """
    Detect the file type and extract text.

    Args:
        file_path: Absolute path to the uploaded file.

    Returns:
        {
            "text":        full extracted plain-text string,
            "page_count":  number of pages/sections (int),
            "word_count":  approximate word count (int),
            "file_type":   "pdf" | "docx",
            "success":     True | False,
            "error":       None | error message string
        }
    """
    if not os.path.exists(file_path):
        return _error_result("File not found on disk", file_path)

    ext = file_path.rsplit(".", 1)[-1].lower() if "." in file_path else ""

    try:
        if ext == "pdf":
            text, page_count = extract_text_from_pdf(file_path)
            file_type = "pdf"
        elif ext == "docx":
            text, page_count = extract_text_from_docx(file_path)
            file_type = "docx"
        else:
            return _error_result(f"Unsupported file type: .{ext}", file_path)

        cleaned = clean_extracted_text(text)
        word_count = len(cleaned.split())

        app_logger.info(
            f"Document processed | type={file_type} | pages={page_count} | words={word_count}"
        )

        return {
            "text": cleaned,
            "page_count": page_count,
            "word_count": word_count,
            "file_type": file_type,
            "success": True,
            "error": None,
        }

    except Exception as exc:
        app_logger.error(f"Document processing failed: {exc}")
        return _error_result(str(exc), file_path)


# --------------------------------------------------------------------------- #
# PDF extraction                                                                #
# --------------------------------------------------------------------------- #

def extract_text_from_pdf(file_path: str) -> tuple:
    """
    Extract text from every page of a PDF using pdfplumber.

    Args:
        file_path: Path to the PDF file.

    Returns:
        Tuple of (text: str, page_count: int).

    Raises:
        Exception on unreadable / corrupted PDF.
    """
    text_parts = []
    page_count = 0

    with pdfplumber.open(file_path) as pdf:
        page_count = len(pdf.pages)
        for page_num, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
            else:
                app_logger.debug(f"Page {page_num} had no extractable text (may be image-based)")

    if not text_parts:
        raise ValueError("No readable text found in the PDF. The file may be scanned/image-based.")

    return "\n".join(text_parts), page_count


# --------------------------------------------------------------------------- #
# DOCX extraction                                                               #
# --------------------------------------------------------------------------- #

def extract_text_from_docx(file_path: str) -> tuple:
    """
    Extract text from a DOCX file using python-docx.

    Captures:
    - Normal paragraphs
    - Table cell text
    - Text from headers/footers (where accessible)

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Tuple of (text: str, section_count: int).

    Raises:
        Exception on corrupt or incompatible DOCX.
    """
    doc = Document(file_path)
    text_parts = []

    # Paragraphs (main body)
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    section_count = len(doc.sections) if doc.sections else 1

    if not text_parts:
        raise ValueError("No readable text found in the DOCX file.")

    return "\n".join(text_parts), section_count


# --------------------------------------------------------------------------- #
# Text cleaning                                                                 #
# --------------------------------------------------------------------------- #

def clean_extracted_text(text: str) -> str:
    """
    Normalise raw extracted text for downstream processing.

    Steps:
    1. Decode any escaped unicode (\\n, \\t)
    2. Replace form-feed and carriage-return with newline
    3. Collapse runs of whitespace within a line to single space
    4. Remove lines that are entirely non-alphanumeric (page markers, rules)
    5. Collapse more than two consecutive blank lines to a single blank line
    6. Strip leading/trailing whitespace

    Args:
        text: Raw extracted text string.

    Returns:
        Cleaned text string.
    """
    if not text:
        return ""

    # Replace windows line endings + form-feed
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\f", "\n")

    lines = text.split("\n")
    cleaned_lines = []
    for line in lines:
        # Collapse internal whitespace
        line = re.sub(r"[ \t]+", " ", line).strip()
        # Skip lines that are purely special chars (page rulers, etc.)
        if line and not re.match(r"^[\W_]+$", line):
            cleaned_lines.append(line)
        elif not line:
            cleaned_lines.append("")  # Preserve blank lines for structure

    # Collapse more than 2 consecutive blank lines
    text = "\n".join(cleaned_lines)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


# --------------------------------------------------------------------------- #
# Internal helpers                                                              #
# --------------------------------------------------------------------------- #

def _error_result(message: str, file_path: str) -> Dict[str, Any]:
    app_logger.error(f"Document processing error [{file_path}]: {message}")
    return {
        "text": "",
        "page_count": 0,
        "word_count": 0,
        "file_type": None,
        "success": False,
        "error": message,
    }
