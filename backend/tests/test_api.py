"""
tests/test_api.py
=================
Integration tests for all Flask API endpoints using the test client.

Run from backend/ directory:
    python -m pytest tests/test_api.py -v
    OR
    python tests/test_api.py
"""

import io
import json
import os
import sys
import unittest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app import create_app


class TestHealthEndpoint(unittest.TestCase):

    def setUp(self):
        self.app = create_app()
        self.app.config["TESTING"] = True
        self.client = self.app.test_client()

    def test_health_returns_200(self):
        resp = self.client.get("/api/health")
        self.assertEqual(resp.status_code, 200)

    def test_health_body(self):
        resp = self.client.get("/api/health")
        data = json.loads(resp.data)
        self.assertEqual(data["status"], "ok")
        self.assertIn("model", data)
        self.assertIn("accuracy", data)

    def test_404_on_unknown_endpoint(self):
        resp = self.client.get("/api/does_not_exist")
        self.assertEqual(resp.status_code, 404)
        data = json.loads(resp.data)
        self.assertFalse(data["success"])


class TestUploadEndpoint(unittest.TestCase):

    def setUp(self):
        self.app = create_app()
        self.app.config["TESTING"] = True
        self.client = self.app.test_client()

    def test_upload_no_file_returns_400(self):
        resp = self.client.post("/api/upload")
        self.assertEqual(resp.status_code, 400)
        data = json.loads(resp.data)
        self.assertFalse(data["success"])

    def test_upload_wrong_type_returns_400(self):
        """Uploading a .txt file should be rejected."""
        data = {"file": (io.BytesIO(b"hello world"), "test.txt")}
        resp = self.client.post("/api/upload", data=data,
                                content_type="multipart/form-data")
        self.assertEqual(resp.status_code, 400)
        body = json.loads(resp.data)
        self.assertFalse(body["success"])

    def test_upload_valid_docx(self):
        """Build a minimal DOCX in memory and upload it."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument()
            doc.add_paragraph("The system shall log all user authentication events.")
            doc.add_paragraph("All passwords must be hashed using bcrypt.")
            doc.add_paragraph("The application should maintain 99.9% uptime.")

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)

            data = {"file": (buf, "test_requirements.docx")}
            resp = self.client.post("/api/upload", data=data,
                                    content_type="multipart/form-data")
            self.assertEqual(resp.status_code, 200)
            body = json.loads(resp.data)
            self.assertTrue(body["success"])
            self.assertIn("file_id", body)
            self.assertEqual(body["file_type"], "docx")
            # file_id available on self for any follow-up assertions in setUp
            self._uploaded_file_id = body["file_id"]
        except ImportError:
            self.skipTest("python-docx not installed")

    def test_upload_empty_filename_returns_400(self):
        data = {"file": (io.BytesIO(b""), "")}
        resp = self.client.post("/api/upload", data=data,
                                content_type="multipart/form-data")
        self.assertIn(resp.status_code, [400, 500])


class TestPredictEndpoint(unittest.TestCase):

    def setUp(self):
        self.app = create_app()
        self.app.config["TESTING"] = True
        self.client = self.app.test_client()

    def test_predict_single_text(self):
        payload = {"text": "The system shall encrypt all user passwords using AES-256."}
        resp = self.client.post("/api/predict",
                                data=json.dumps(payload),
                                content_type="application/json")
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        self.assertTrue(data["success"])
        self.assertEqual(data["count"], 1)
        self.assertIn("category", data["results"][0])

    def test_predict_batch_texts(self):
        payload = {"texts": [
            "The system must respond in under 2 seconds.",
            "Users shall be able to reset their password.",
            "All data must be backed up every 24 hours.",
        ]}
        resp = self.client.post("/api/predict",
                                data=json.dumps(payload),
                                content_type="application/json")
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        self.assertEqual(data["count"], 3)

    def test_predict_missing_body_returns_400(self):
        resp = self.client.post("/api/predict")
        self.assertEqual(resp.status_code, 400)

    def test_predict_empty_texts_list_returns_400(self):
        payload = {"texts": []}
        resp = self.client.post("/api/predict",
                                data=json.dumps(payload),
                                content_type="application/json")
        self.assertEqual(resp.status_code, 400)

    def test_predict_confidence_in_range(self):
        payload = {"text": "The system shall authenticate users before granting access."}
        resp = self.client.post("/api/predict",
                                data=json.dumps(payload),
                                content_type="application/json")
        data = json.loads(resp.data)
        confidence = data["results"][0]["confidence"]
        self.assertGreaterEqual(confidence, 0)
        self.assertLessEqual(confidence, 100)


class TestReportEndpoint(unittest.TestCase):

    def setUp(self):
        self.app = create_app()
        self.app.config["TESTING"] = True
        self.client = self.app.test_client()

    def test_report_unknown_id_returns_404(self):
        resp = self.client.get("/api/report/nonexistent_id_xyz")
        self.assertEqual(resp.status_code, 404)
        data = json.loads(resp.data)
        self.assertFalse(data["success"])


class TestAnalyzeEndpoint(unittest.TestCase):

    def setUp(self):
        self.app = create_app()
        self.app.config["TESTING"] = True
        self.client = self.app.test_client()

    def test_analyze_missing_file_id_returns_400(self):
        resp = self.client.post("/api/analyze",
                                data=json.dumps({}),
                                content_type="application/json")
        self.assertEqual(resp.status_code, 400)

    def test_analyze_unknown_file_id_returns_404(self):
        resp = self.client.post("/api/analyze",
                                data=json.dumps({"file_id": "unknown_xyz_000"}),
                                content_type="application/json")
        self.assertEqual(resp.status_code, 404)

    def test_full_pipeline_docx(self):
        """Upload a DOCX then analyze it end-to-end."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument()
            requirements = [
                "The system shall authenticate all users using a secure login mechanism.",
                "All passwords must be stored using bcrypt hashing with a salt.",
                "The application shall support a minimum of 500 concurrent sessions.",
                "The system must respond to all API requests within 3 seconds.",
                "Users shall be able to upload documents up to 10 MB in size.",
                "The system should provide an audit log of all administrative actions.",
                "The application shall be deployable on Linux and Windows platforms.",
                "All data transmissions shall be encrypted using TLS 1.3.",
            ]
            for req in requirements:
                doc.add_paragraph(req)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)

            # 1. Upload
            upload_resp = self.client.post(
                "/api/upload",
                data={"file": (buf, "e2e_test.docx")},
                content_type="multipart/form-data",
            )
            self.assertEqual(upload_resp.status_code, 200, upload_resp.data)
            file_id = json.loads(upload_resp.data)["file_id"]

            # 2. Analyze
            analyze_resp = self.client.post(
                "/api/analyze",
                data=json.dumps({"file_id": file_id}),
                content_type="application/json",
            )
            self.assertEqual(analyze_resp.status_code, 200, analyze_resp.data)
            analysis = json.loads(analyze_resp.data)
            self.assertTrue(analysis["success"])
            self.assertGreater(analysis["total_requirements"], 0)
            self.assertIn("overall_score", analysis)
            self.assertIn("category_scores", analysis)
            self.assertIn("recommendations", analysis)

            # 3. Fetch report
            report_resp = self.client.get(f"/api/report/{file_id}")
            self.assertEqual(report_resp.status_code, 200)
            report = json.loads(report_resp.data)
            self.assertTrue(report["success"])
            self.assertIn("summary", report)
            self.assertIn("requirements", report)

        except ImportError:
            self.skipTest("python-docx not installed")


class TestAnalysesListEndpoint(unittest.TestCase):

    def setUp(self):
        self.app = create_app()
        self.app.config["TESTING"] = True
        self.client = self.app.test_client()

    def test_analyses_returns_200(self):
        resp = self.client.get("/api/analyses")
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        self.assertTrue(data["success"])
        self.assertIn("analyses", data)
        self.assertIsInstance(data["analyses"], list)


if __name__ == "__main__":
    unittest.main(verbosity=2)
