"""
tests/test_document_processor.py
==================================
Unit tests for services/document_processor.py

Run from backend/ directory:
    python -m pytest tests/test_document_processor.py -v
    OR
    python tests/test_document_processor.py
"""

import os
import sys
import unittest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from services.document_processor import (
    clean_extracted_text,
    process_document,
)


class TestCleanExtractedText(unittest.TestCase):
    """Tests for the text cleaning function — no file I/O needed."""

    def test_strips_leading_trailing_whitespace(self):
        result = clean_extracted_text("  hello world  ")
        self.assertEqual(result, "hello world")

    def test_collapses_internal_whitespace(self):
        result = clean_extracted_text("hello   world\there")
        self.assertNotIn("   ", result)
        self.assertNotIn("\t", result)

    def test_removes_form_feed(self):
        result = clean_extracted_text("page1\fpage2")
        self.assertNotIn("\f", result)

    def test_collapses_excessive_blank_lines(self):
        text = "line1\n\n\n\n\nline2"
        result = clean_extracted_text(text)
        self.assertNotIn("\n\n\n", result)

    def test_empty_string(self):
        self.assertEqual(clean_extracted_text(""), "")

    def test_none_like_empty(self):
        self.assertEqual(clean_extracted_text(""), "")

    def test_only_special_chars_removed(self):
        """Lines that are purely special chars should be dropped."""
        text = "header\n--------------------\nbody text here"
        result = clean_extracted_text(text)
        self.assertNotIn("----", result)

    def test_preserves_requirement_sentences(self):
        text = "The system shall encrypt all passwords. The app must respond in 2 seconds."
        result = clean_extracted_text(text)
        self.assertIn("shall encrypt", result)
        self.assertIn("must respond", result)


class TestProcessDocument(unittest.TestCase):
    """Tests for process_document() — tests against non-existent files
    to verify error handling without needing actual PDFs."""

    def test_missing_file_returns_failure(self):
        result = process_document("/nonexistent/path/file.pdf")
        self.assertFalse(result["success"])
        self.assertIsNotNone(result["error"])
        self.assertEqual(result["text"], "")

    def test_unsupported_extension_returns_failure(self):
        # Create a temp txt file to avoid "not found" error
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as f:
            f.write(b"some text")
            tmp_path = f.name
        try:
            result = process_document(tmp_path)
            self.assertFalse(result["success"])
            self.assertIn("Unsupported", result["error"])
        finally:
            os.unlink(tmp_path)

    def test_result_keys_always_present(self):
        """Even on failure, all keys must be present in the result."""
        result = process_document("/no/such/file.pdf")
        for key in ("text", "page_count", "word_count", "file_type", "success", "error"):
            self.assertIn(key, result)

    def test_process_docx_from_scratch(self):
        """Create a minimal DOCX in memory and verify extraction."""
        try:
            from docx import Document as DocxDocument
            import tempfile

            doc = DocxDocument()
            doc.add_paragraph("The system shall authenticate all users before granting access.")
            doc.add_paragraph("All data must be encrypted at rest using AES-256.")

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                doc.save(f.name)
                tmp_path = f.name

            result = process_document(tmp_path)
            self.assertTrue(result["success"], msg=result.get("error"))
            self.assertIn("authenticate", result["text"])
            self.assertGreater(result["word_count"], 5)
        except ImportError:
            self.skipTest("python-docx not installed")
        finally:
            if "tmp_path" in dir():
                os.unlink(tmp_path)


if __name__ == "__main__":
    unittest.main(verbosity=2)
