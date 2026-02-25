"""
tests/test_e2e_workflow.py
==========================
End-to-end workflow tests for the complete QualityMapAI pipeline.

Covers
------
- 3.1  Upload → Analyze → Report full flow (multiple file types)
- 3.1  Large file stress test (many requirements)
- 3.1  Re-analysis (analyze the same file twice)
- 3.1  Dashboard reflects new analyses
- 3.2  Upload failures & malformed files
- 3.2  No requirements detected (empty / non-requirement text)
- 3.2  Predict endpoint edge cases

Run:
    cd backend
    python -m pytest tests/test_e2e_workflow.py -v
"""

import io
import json
import os
import sys
import unittest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app import create_app


# ──────────────────────────────────────────────────────────────────────────── #
# HELPERS                                                                       #
# ──────────────────────────────────────────────────────────────────────────── #

def _make_docx(paragraphs: list[str]) -> io.BytesIO:
    """Create a minimal DOCX in memory with the given paragraphs."""
    from docx import Document as DocxDocument
    doc = DocxDocument()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


SAMPLE_REQUIREMENTS = [
    "The system shall authenticate all users using a secure login mechanism.",
    "All passwords must be stored using bcrypt hashing with a salt.",
    "The application shall support a minimum of 500 concurrent sessions.",
    "The system must respond to all API requests within 3 seconds.",
    "Users shall be able to upload documents up to 10 MB in size.",
    "The system should provide an audit log of all administrative actions.",
    "The application shall be deployable on Linux and Windows platforms.",
    "All data transmissions shall be encrypted using TLS 1.3.",
    "The system shall provide role-based access control for administrators.",
    "Error messages must not expose internal system information.",
    "The application should support internationalisation for at least 3 languages.",
    "The database must be backed up automatically every 24 hours.",
]


# ──────────────────────────────────────────────────────────────────────────── #
# E2E WORKFLOW: Upload → Analyze → Report                                       #
# ──────────────────────────────────────────────────────────────────────────── #

class TestE2EWorkflow(unittest.TestCase):
    """Full pipeline: upload → analyze → report → dashboard."""

    @classmethod
    def setUpClass(cls):
        cls.app = create_app()
        cls.app.config["TESTING"] = True
        cls.client = cls.app.test_client()

    # ── helper ──────────────────────────────────────────────────────────── #
    def _upload_docx(self, paragraphs, filename="e2e_test.docx"):
        buf = _make_docx(paragraphs)
        resp = self.client.post(
            "/api/upload",
            data={"file": (buf, filename)},
            content_type="multipart/form-data",
        )
        return resp

    # ── tests ───────────────────────────────────────────────────────────── #
    def test_01_full_pipeline_upload_analyze_report(self):
        """Upload a DOCX → Analyze → Fetch Report. Verify every key field."""
        # 1. Upload
        resp = self._upload_docx(SAMPLE_REQUIREMENTS)
        self.assertEqual(resp.status_code, 200, f"Upload failed: {resp.data}")
        upload_data = json.loads(resp.data)
        self.assertTrue(upload_data["success"])
        file_id = upload_data["file_id"]
        self.assertTrue(len(file_id) > 0)
        self.assertEqual(upload_data["file_type"], "docx")

        # 2. Analyze
        resp = self.client.post(
            "/api/analyze",
            data=json.dumps({"file_id": file_id}),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 200, f"Analyze failed: {resp.data}")
        analysis = json.loads(resp.data)
        self.assertTrue(analysis["success"])
        self.assertGreater(analysis["total_requirements"], 0)
        self.assertIn("domain", analysis)
        self.assertIn("domain", analysis["domain"])
        self.assertIsInstance(analysis["category_scores"], dict)
        self.assertIsInstance(analysis["requirements"], list)
        self.assertIsInstance(analysis["recommendations"], list)
        self.assertIsInstance(analysis["gap_analysis"], list)
        self.assertIsInstance(analysis["categories_present"], list)
        self.assertIsInstance(analysis["categories_missing"], list)
        self.assertIsInstance(analysis["processing_time_s"], (int, float))

        # 3. Verify individual requirement structure
        for req in analysis["requirements"]:
            self.assertIn("text", req)
            self.assertIn("category", req)
            self.assertIn("confidence", req)
            self.assertGreaterEqual(req["confidence"], 0)
            self.assertLessEqual(req["confidence"], 100.0)

        # 4. Fetch Report
        resp = self.client.get(f"/api/report/{file_id}")
        self.assertEqual(resp.status_code, 200, f"Report failed: {resp.data}")
        report = json.loads(resp.data)
        self.assertTrue(report["success"])
        self.assertEqual(report["analysis_id"], file_id)
        self.assertIn("summary", report)
        self.assertIn("domain", report["summary"])
        self.assertIn("total_requirements", report["summary"])
        self.assertEqual(
            report["summary"]["total_requirements"],
            analysis["total_requirements"],
        )
        self.assertIn("requirements", report)
        self.assertEqual(len(report["requirements"]), analysis["total_requirements"])

    def test_02_re_analysis_overwrites(self):
        """Analyzing the same file twice should overwrite (not duplicate)."""
        resp = self._upload_docx(SAMPLE_REQUIREMENTS, "reanalysis_test.docx")
        file_id = json.loads(resp.data)["file_id"]

        # First analysis
        resp1 = self.client.post(
            "/api/analyze",
            data=json.dumps({"file_id": file_id}),
            content_type="application/json",
        )
        self.assertEqual(resp1.status_code, 200)

        # Second analysis (should not error — re-analysis supported)
        resp2 = self.client.post(
            "/api/analyze",
            data=json.dumps({"file_id": file_id}),
            content_type="application/json",
        )
        self.assertEqual(resp2.status_code, 200)
        data2 = json.loads(resp2.data)
        self.assertTrue(data2["success"])

        # Report should still work
        resp3 = self.client.get(f"/api/report/{file_id}")
        self.assertEqual(resp3.status_code, 200)

    def test_03_dashboard_reflects_analyses(self):
        """After uploading & analyzing, the analyses list should contain the entry."""
        resp = self._upload_docx(SAMPLE_REQUIREMENTS[:4], "dashboard_test.docx")
        file_id = json.loads(resp.data)["file_id"]

        self.client.post(
            "/api/analyze",
            data=json.dumps({"file_id": file_id}),
            content_type="application/json",
        )

        resp = self.client.get("/api/analyses")
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        self.assertTrue(data["success"])
        ids = [a["analysis_id"] for a in data["analyses"]]
        self.assertIn(file_id, ids)

    def test_04_category_scores_have_required_keys(self):
        """Each category score dict should have count, percentage, etc."""
        resp = self._upload_docx(SAMPLE_REQUIREMENTS, "category_test.docx")
        file_id = json.loads(resp.data)["file_id"]

        resp = self.client.post(
            "/api/analyze",
            data=json.dumps({"file_id": file_id}),
            content_type="application/json",
        )
        analysis = json.loads(resp.data)
        for cat_name, cat_data in analysis["category_scores"].items():
            self.assertIn("count", cat_data)
            self.assertIn("percentage", cat_data)
            self.assertIsInstance(cat_data["count"], int)
            self.assertGreaterEqual(cat_data["count"], 0)


# ──────────────────────────────────────────────────────────────────────────── #
# STRESS TEST: Large file with many requirements                                #
# ──────────────────────────────────────────────────────────────────────────── #

class TestStressLargeFile(unittest.TestCase):
    """Upload a file with many requirements and verify the pipeline handles it."""

    @classmethod
    def setUpClass(cls):
        cls.app = create_app()
        cls.app.config["TESTING"] = True
        cls.client = cls.app.test_client()

    def test_large_document_50_requirements(self):
        """A DOCX with 50 requirement sentences should be handled without error."""
        base_reqs = [
            "The system shall support user authentication via OAuth 2.0.",
            "All API endpoints must validate input parameters.",
            "The database should handle at least 1000 concurrent connections.",
            "Logs must be retained for a minimum of 90 days.",
            "The application shall provide CSV export of analysis results.",
        ]
        # Repeat to reach 50
        large_reqs = (base_reqs * 10)[:50]
        buf = _make_docx(large_reqs)

        # Upload
        resp = self.client.post(
            "/api/upload",
            data={"file": (buf, "stress_large.docx")},
            content_type="multipart/form-data",
        )
        self.assertEqual(resp.status_code, 200)
        file_id = json.loads(resp.data)["file_id"]

        # Analyze
        resp = self.client.post(
            "/api/analyze",
            data=json.dumps({"file_id": file_id}),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        self.assertTrue(data["success"])
        self.assertGreater(data["total_requirements"], 0)
        self.assertIsInstance(data["processing_time_s"], (int, float))

    def test_large_document_100_requirements(self):
        """A DOCX with 100 requirement sentences should still succeed."""
        templates = [
            "The system shall provide feature {} for compliance.",
            "Users must be able to perform action {} within 5 seconds.",
            "The application should log event {} for audit purposes.",
            "All modules shall support configuration option {}.",
            "The system must validate input {} before processing.",
        ]
        large_reqs = [t.format(i) for i in range(100) for t in templates][:100]
        buf = _make_docx(large_reqs)

        resp = self.client.post(
            "/api/upload",
            data={"file": (buf, "stress_100.docx")},
            content_type="multipart/form-data",
        )
        self.assertEqual(resp.status_code, 200)
        file_id = json.loads(resp.data)["file_id"]

        resp = self.client.post(
            "/api/analyze",
            data=json.dumps({"file_id": file_id}),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        self.assertTrue(data["success"])
        self.assertGreaterEqual(data["total_requirements"], 20)


# ──────────────────────────────────────────────────────────────────────────── #
# ERROR SCENARIOS                                                               #
# ──────────────────────────────────────────────────────────────────────────── #

class TestErrorScenarios(unittest.TestCase):
    """Verify the system handles error conditions gracefully."""

    @classmethod
    def setUpClass(cls):
        cls.app = create_app()
        cls.app.config["TESTING"] = True
        cls.client = cls.app.test_client()

    # ── Upload failures ─────────────────────────────────────────────────── #
    def test_upload_no_file_field(self):
        resp = self.client.post("/api/upload")
        self.assertEqual(resp.status_code, 400)
        data = json.loads(resp.data)
        self.assertFalse(data["success"])
        self.assertIn("error", data)

    def test_upload_unsupported_file_type(self):
        resp = self.client.post(
            "/api/upload",
            data={"file": (io.BytesIO(b"text content"), "notes.txt")},
            content_type="multipart/form-data",
        )
        self.assertEqual(resp.status_code, 400)

    def test_upload_empty_filename(self):
        resp = self.client.post(
            "/api/upload",
            data={"file": (io.BytesIO(b""), "")},
            content_type="multipart/form-data",
        )
        self.assertIn(resp.status_code, [400, 500])

    def test_upload_exe_file_rejected(self):
        resp = self.client.post(
            "/api/upload",
            data={"file": (io.BytesIO(b"\x4d\x5a" + b"\x00" * 100), "malware.exe")},
            content_type="multipart/form-data",
        )
        self.assertEqual(resp.status_code, 400)

    # ── Analyze failures ────────────────────────────────────────────────── #
    def test_analyze_missing_json_body(self):
        resp = self.client.post("/api/analyze")
        self.assertEqual(resp.status_code, 400)

    def test_analyze_empty_file_id(self):
        resp = self.client.post(
            "/api/analyze",
            data=json.dumps({"file_id": ""}),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 400)

    def test_analyze_nonexistent_file_id(self):
        resp = self.client.post(
            "/api/analyze",
            data=json.dumps({"file_id": "does_not_exist_12345"}),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 404)

    def test_analyze_whitespace_file_id(self):
        resp = self.client.post(
            "/api/analyze",
            data=json.dumps({"file_id": "   "}),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 400)

    # ── No requirements detected ────────────────────────────────────────── #
    def test_analyze_no_requirements_found(self):
        """Upload a DOCX with no requirement-like sentences → 422."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            self.skipTest("python-docx not installed")

        doc = DocxDocument()
        doc.add_paragraph("Hello world.")
        doc.add_paragraph("The cat sat on the mat.")
        doc.add_paragraph("It rained.")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        # Upload
        resp = self.client.post(
            "/api/upload",
            data={"file": (buf, "no_reqs.docx")},
            content_type="multipart/form-data",
        )
        self.assertEqual(resp.status_code, 200)
        file_id = json.loads(resp.data)["file_id"]

        # Analyze – should return 422 (no requirements)
        resp = self.client.post(
            "/api/analyze",
            data=json.dumps({"file_id": file_id}),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 422)
        data = json.loads(resp.data)
        self.assertFalse(data["success"])

    # ── Report failures ─────────────────────────────────────────────────── #
    def test_report_nonexistent_id(self):
        resp = self.client.get("/api/report/nonexistent_xyz")
        self.assertEqual(resp.status_code, 404)
        data = json.loads(resp.data)
        self.assertFalse(data["success"])

    # ── Predict failures ────────────────────────────────────────────────── #
    def test_predict_no_body(self):
        resp = self.client.post("/api/predict")
        self.assertEqual(resp.status_code, 400)

    def test_predict_empty_text(self):
        resp = self.client.post(
            "/api/predict",
            data=json.dumps({"text": ""}),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 400)

    def test_predict_whitespace_text(self):
        resp = self.client.post(
            "/api/predict",
            data=json.dumps({"text": "   "}),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 400)

    def test_predict_non_string_text(self):
        resp = self.client.post(
            "/api/predict",
            data=json.dumps({"text": 12345}),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 400)

    def test_predict_texts_with_empty_entry(self):
        resp = self.client.post(
            "/api/predict",
            data=json.dumps({"texts": ["Valid requirement.", ""]}),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 400)

    def test_predict_missing_key(self):
        resp = self.client.post(
            "/api/predict",
            data=json.dumps({"wrong_key": "value"}),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 400)

    # ── HTTP method errors ──────────────────────────────────────────────── #
    def test_get_on_upload_returns_405(self):
        resp = self.client.get("/api/upload")
        self.assertEqual(resp.status_code, 405)

    def test_get_on_analyze_returns_405(self):
        resp = self.client.get("/api/analyze")
        self.assertEqual(resp.status_code, 405)

    def test_get_on_predict_returns_405(self):
        resp = self.client.get("/api/predict")
        self.assertEqual(resp.status_code, 405)

    def test_post_on_health_returns_405(self):
        resp = self.client.post("/api/health")
        self.assertEqual(resp.status_code, 405)


# ──────────────────────────────────────────────────────────────────────────── #
# PREDICT EDGE CASES                                                            #
# ──────────────────────────────────────────────────────────────────────────── #

class TestPredictEdgeCases(unittest.TestCase):
    """Extra coverage for the /predict endpoint."""

    @classmethod
    def setUpClass(cls):
        cls.app = create_app()
        cls.app.config["TESTING"] = True
        cls.client = cls.app.test_client()

    def test_predict_very_long_text(self):
        """A very long requirement should still classify without crashing."""
        long_text = "The system shall " + "handle data " * 500 + "securely."
        resp = self.client.post(
            "/api/predict",
            data=json.dumps({"text": long_text}),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        self.assertTrue(data["success"])

    def test_predict_special_characters(self):
        """Text with special chars should not crash the classifier."""
        resp = self.client.post(
            "/api/predict",
            data=json.dumps({"text": "The system shall handle <html> & \"quotes\" safely."}),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 200)

    def test_predict_unicode_text(self):
        resp = self.client.post(
            "/api/predict",
            data=json.dumps({"text": "システムはユーザー認証を提供しなければならない。"}),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 200)

    def test_predict_large_batch(self):
        """Batch of 50 texts should succeed."""
        texts = [f"The system shall support feature number {i}." for i in range(50)]
        resp = self.client.post(
            "/api/predict",
            data=json.dumps({"texts": texts}),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        self.assertEqual(data["count"], 50)


if __name__ == "__main__":
    unittest.main(verbosity=2)
